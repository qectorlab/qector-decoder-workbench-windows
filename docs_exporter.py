"""docs_exporter.py  -  Robust in-app export of the official documentation set.

One entry point for the "Export Official Docs" GUI button (and any other
caller): ``export_public_docs(outdir, on_log=None)``.  It drives the same two
generators the release pipeline uses (``generate_manuals.build_public_docs``
and ``api_reference.build_api_reference``) so in-app exports are byte-identical
to the shipped release docs.  This module never raises: every failure is
recorded in the returned per-artifact report and, when a logger callback is
supplied, surfaced through it.
"""

from __future__ import annotations

import traceback
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Optional

from api_reference import build_api_reference
from generate_manuals import build_public_docs
from utils import get_export_dir, sha256_of, write_sha256_manifest


def default_docs_dir() -> Path:
    """Per-user directory for official doc exports (created on demand)."""
    return get_export_dir() / "official_docs"


def docx_available() -> bool:
    """True when python-docx is importable, so DOCX export can be attempted."""
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def export_to_docx(md_path: Path, docx_path: Path) -> tuple[bool, str]:
    """Convert a Markdown document to DOCX.  Returns ``(ok, message)``.

    There is deliberately no "write the raw text and call it a .docx" fallback:
    a file with a .docx extension that Word cannot open is worse than a missing
    file, because the failure only surfaces when someone tries to open it.  If
    python-docx is unavailable this reports that plainly and writes nothing.
    """
    try:
        import docx
    except Exception:
        return False, ("python-docx is not installed, so DOCX export was skipped "
                       "(pip install python-docx). The Markdown and PDF versions "
                       "are unaffected.")
    try:
        document = docx.Document()
        in_code_block = False
        for raw_line in md_path.read_text(encoding="utf-8").splitlines():
            line = raw_line.strip()
            if line.startswith("```"):
                in_code_block = not in_code_block
                continue
            if in_code_block:
                if line:
                    para = document.add_paragraph(line)
                    para.style = document.styles["No Spacing"]
                continue
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                document.add_heading(line[level:].strip(), level=min(level, 4))
            elif line.startswith(("- ", "* ")):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|"):
                # Markdown tables round-trip poorly; keep the row readable
                # rather than dropping the data.
                document.add_paragraph(line.strip("|").replace("|", "   "))
            elif line:
                document.add_paragraph(line)
        document.save(str(docx_path))
        return True, f"DOCX written to {docx_path}"
    except Exception as exc:
        return False, f"DOCX export failed: {type(exc).__name__}: {exc}"


def copy_to_clipboard(text: str, widget: Any = None) -> tuple[bool, str]:
    """Copy ``text`` to the system clipboard.  Returns ``(ok, message)``.

    ``widget`` must be any live Tk widget from the running application.  The
    clipboard is owned by a Tk interpreter, so this deliberately reuses the
    app's existing root instead of creating a second one: two live Tk
    interpreters in one process is a known source of flakiness, and a root that
    is destroyed straight after the copy hands ownership back before the paste
    ever happens.
    """
    if widget is None:
        return False, ("clipboard copy needs a live Tk widget from the running app; "
                       "no widget was supplied")
    try:
        widget.clipboard_clear()
        widget.clipboard_append(text)
        widget.update_idletasks()
        return True, f"Copied {len(text)} characters to the clipboard"
    except Exception as exc:
        return False, f"Clipboard copy failed: {type(exc).__name__}: {exc}"


def export_public_docs(
    outdir: Optional[Path] = None,
    on_log: Optional[Callable[[str], None]] = None,
    include_api_reference: bool = True,
    include_docx: bool = True,
) -> dict[str, dict]:
    """Generate the full public documentation set into ``outdir``.

    Parameters
    ----------
    outdir:
        Target directory.  Defaults to the per-user ``exports/official_docs``.
    on_log:
        Optional ``on_log(message)`` callback invoked for each completed or
        failed artifact (thread-safe use is the caller's responsibility).
    include_api_reference:
        Also generate ``QECTOR_API_Reference.md`` and ``.pdf``.
    include_docx:
        Generate DOCX versions of key Markdown manuals.

    Returns
    -------
    ``{name: {"ok": bool, "path": Path|None, "error": str|None}}`` for every
    artifact attempted, plus a final ``"_summary"`` entry.  Never raises.
    """
    if outdir is None:
        outdir = default_docs_dir()
    outdir = Path(outdir)
    report: dict[str, dict] = {}
    failures: list[str] = []

    def log(msg: str) -> None:
        if on_log is not None:
            try:
                on_log(msg)
            except Exception:
                pass

    try:
        outdir.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        return {
            "_summary": {
                "ok": False,
                "path": None,
                "error": f"could not create export directory {outdir}: {exc}",
            }
        }

    started = datetime.now(timezone.utc)
    log(f"Official docs export started at {started.isoformat(timespec='seconds')}Z")

    # 1) Manuals set (user manuals, quick start, MCP guide, LLM json, README, zip)
    try:
        produced = build_public_docs(outdir)
        for name, path in produced.items():
            try:
                size = path.stat().st_size
                try:
                    digest = sha256_of(path)
                except Exception:
                    digest = None
                report[name] = {"ok": True, "path": path, "error": None,
                                "bytes": size, "sha256": digest}
                log(f"  [OK] {name} ({size} bytes)"
                    + ("" if digest else "  [sha256 unavailable]"))

                # DOCX companion for each Markdown manual.  A missing
                # python-docx is reported once, not silently papered over.
                if include_docx and path.suffix.lower() == ".md":
                    docx_name = path.stem + ".docx"
                    docx_path = outdir / docx_name
                    docx_ok, docx_msg = export_to_docx(path, docx_path)
                    if docx_ok:
                        docx_size = docx_path.stat().st_size
                        try:
                            docx_digest = sha256_of(docx_path)
                        except Exception:
                            docx_digest = None
                        report[docx_name] = {"ok": True, "path": docx_path,
                                             "error": None, "bytes": docx_size,
                                             "sha256": docx_digest}
                        log(f"  [OK] {docx_name} ({docx_size} bytes)")
                    else:
                        report[docx_name] = {"ok": False, "path": None,
                                             "error": docx_msg, "bytes": 0}
                        log(f"  [SKIP] {docx_name}: {docx_msg}")
            except Exception as e:
                report[name] = {"ok": False, "path": path, "error": str(e), "bytes": 0}
    except Exception as exc:
        msg = f"manuals set failed: {exc}\n{traceback.format_exc()}"
        failures.append(msg)
        report["manuals"] = {"ok": False, "path": None, "error": str(exc)}
        log(f"  [FAIL] manuals set: {exc}")

    # 2) API reference (md + pdf)
    if include_api_reference:
        try:
            api = build_api_reference(outdir)
            for name, path in api.items():
                try:
                    size = path.stat().st_size
                except Exception:
                    size = -1
                try:
                    digest = sha256_of(path)
                except Exception:
                    digest = None
                report[name] = {"ok": True, "path": path, "error": None,
                                "bytes": size, "sha256": digest}
                log(f"  [OK] {name} ({size} bytes)")
        except Exception as exc:
            msg = f"API reference failed: {exc}\n{traceback.format_exc()}"
            failures.append(msg)
            report["API reference"] = {"ok": False, "path": None, "error": str(exc)}
            log(f"  [FAIL] API reference: {exc}")

    # A real SHA-256 manifest for the export: every artifact this run produced,
    # in one coreutils file so the whole set verifies with a single
    # ``sha256sum -c``.
    ok_paths = [v["path"] for v in report.values()
                if v.get("ok") and v.get("path") and Path(v["path"]).is_file()]
    sum_ok, sum_msg = write_sha256_manifest(outdir, ok_paths)
    if sum_ok:
        manifest_path = Path(sum_msg)
        try:
            manifest_digest = sha256_of(manifest_path)
        except Exception:
            manifest_digest = None
        report["SHA256SUMS.txt"] = {
            "ok": True, "path": manifest_path, "error": None,
            "bytes": manifest_path.stat().st_size,
            "sha256": manifest_digest,
        }
        log(f"  [OK] SHA256SUMS.txt ({len(ok_paths)} artifacts checksummed)")
    else:
        log(f"  [WARN] SHA256SUMS.txt not written: {sum_msg}")

    done = datetime.now(timezone.utc)
    elapsed = (done - started).total_seconds()
    ok_count = sum(1 for v in report.values() if v.get("ok"))
    report["_summary"] = {
        "ok": not failures,
        "path": outdir,
        "error": "; ".join(failures) if failures else None,
        "artifacts_ok": ok_count,
        "artifacts_total": len(report),
        "elapsed_s": round(elapsed, 2),
    }
    if failures:
        log(f"Export finished with {len(failures)} failure(s) in {elapsed:.1f}s")
    else:
        log(f"Export complete: {ok_count} artifacts in {elapsed:.1f}s")
    return report

